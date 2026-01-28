"""
ACQUA Report Reviewer - Process and analyze ACQUA test reports
===============================================================

A tool for processing ACQUA Word document reports and extracting
test data, validation results, and status information.

Author: Jian Zou
"""

__version__ = "1.0.0"
__author__ = "Jian Zou"
__email__ = "jianzou@microsoft.com"
__date__ = "2026-01-27"
__description__ = "ACQUA Report Reviewer - Process and analyze ACQUA test reports"

import tkinter as tk
from tkinter import filedialog
import csv
import re
import os
import sys
from datetime import datetime

def get_version_info():
    """Return version information as a formatted string."""
    return f"""
╔══════════════════════════════════════════════════════════════╗
║  ACQUA Report Reviewer v{__version__:<37}║
║  Author: {__author__:<50}║
║  Last Updated: {__date__:<44}║
╚══════════════════════════════════════════════════════════════╝
"""

# Try to import python-docx. If not installed, inform the user.
try:
    from docx import Document
except ImportError:
    print("Error: The 'python-docx' library is missing.")
    print("Please install it using the command: pip install python-docx")
    input("Press Enter to exit...")
    sys.exit(1)

def extract_clean_date(text):
    """
    Extracts the date and time from the SmdDate style text.
    Format expected: 6/19/2025 4:14 PM, followed by ACQUA
    Strictly filters for lines containing 'ACQUA'.
    """
    # Regex look for Date + Time ... ACQUA
    pattern = r"(\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s+(?:AM|PM)).*?ACQUA"
    
    match = re.search(pattern, text)
    if match:
        return match.group(1).strip()
    
    return None

def extract_acqua_database_info(file_paths):
    """
    Extract ACQUA version and Database version from each provided file.
    Searches for patterns like "ACQUA 6.0.200" and "Database Version: 51_MS_Teams_Rev05_SP2".
    Returns list of dicts with file info, ACQUA version, and database version.
    """
    results = []
    
    # Regex patterns - ACQUA X.X.X where X can be 1-3 digits
    acqua_pattern = r"ACQUA\s+(\d{1,3}\.\d{1,3}\.\d{1,3})"
    database_pattern = r"Database\s+Version:\s*([^\n\r]+)"
    
    for file_path in file_paths:
        try:
            doc = Document(file_path)
            file_name = os.path.basename(file_path)
            
            acqua_version = None
            database_version = None
            
            # Search through all paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Find first ACQUA version
                if not acqua_version:
                    acqua_match = re.search(acqua_pattern, text)
                    if acqua_match:
                        acqua_version = f"ACQUA {acqua_match.group(1)}"
                
                # Find first Database Version
                if not database_version:
                    db_match = re.search(database_pattern, text, re.IGNORECASE)
                    if db_match:
                        database_version = db_match.group(1).strip()
                
                # Stop if both found
                if acqua_version and database_version:
                    break
            
            # Also search through tables if not found in paragraphs
            if not acqua_version or not database_version:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            
                            if not acqua_version:
                                acqua_match = re.search(acqua_pattern, cell_text)
                                if acqua_match:
                                    acqua_version = f"ACQUA {acqua_match.group(1)}"
                            
                            if not database_version:
                                db_match = re.search(database_pattern, cell_text, re.IGNORECASE)
                                if db_match:
                                    database_version = db_match.group(1).strip()
                            
                            if acqua_version and database_version:
                                break
                        if acqua_version and database_version:
                            break
                    if acqua_version and database_version:
                        break
            
            results.append({
                'file': file_name,
                'file_path': file_path,
                'acqua_version': acqua_version if acqua_version else 'Not Found',
                'database_version': database_version if database_version else 'Not Found'
            })
            
        except Exception as e:
            results.append({
                'file': os.path.basename(file_path),
                'file_path': file_path,
                'acqua_version': f'Error: {str(e)}',
                'database_version': f'Error: {str(e)}'
            })
    
    return results

# Define Code Categories and Lists
CODE_DEFINITIONS = {
    "P-series (A)": [
        "P01D", "P01A", "P02A", "P03A", "P04A", "P05A", "P06A", "P07A", "P08A", "P09A", "P10A", 
        "P11A", "P12A", "P13A", "P14A", "P15A", "P16A", "P17A", "P18A", "P19A", "P20A", 
        "P21A", "P22A", "P23A", "P24A", "P25A", "P26A", "P27A", "P28A", "P29A", "P30A"
    ],
    "P-series (R)": [
        "P01R", "P02R", "P03R", "P04R", "P05R", "P06R", "P07R", "P08R", "P09R", "P10R", "P11R", "P12R"
    ],
    "Device-direct (Di)": [
        "Di01A", "Di02A", "Di03A", "Di04A", "Di05A", "Di06A", "Di07A",  "Di09A", "Di10A", "Di11A", "Di12A"
    ],
    "Option codes": [
        "OpO01R", "OpO02R", "OpO03R", "OpO04R", "OpO05R", 
        "OpM01", "OpM02", "OpM03", "OpM04", "OpM05", "OpM06", 
        "OpS01A", "OpS02A", "OpS03A", "OpS04A", "OpS05A", "OpS06A", "OpS07A", "OpS08A", "OpS09A", "OpS10A", 
        "OpS11A", "OpS12A", "OpS13A", "OpS14A", "OpS15A", "OpS16A", "OpS17A", "OpS18A", "OpS19A", "OpS20A", 
        "OpS01R", "OpS02R", "OpS03R", "OpS04R", "OpS05R", 
        "OpP01A", "OpP02A", "OpP03A"
    ]
}

REQUIRED_SHARED_SPEAKERPHONE = {
    "AR": {
        "P01A", "P02A", "P07A", "P08A", "P09A", "P10A", "P11A", "P13A", "P14A", "P15A", 
        "P16A", "P17A", "P18A", "P19A", "P20A", "P21A", "P24A", "P25A", "P26A"
    },
    "RR": {
        "P01R", "P02R", "P03R", "P04R", "P05R", "P10R", "P06R", "P07R", "P08R", "P09R", "P11R", "P12R"
    },
    "Di": {
        "Di01A", "Di03A", "Di05A", "Di06A", "Di07A",  "Di09A", "Di10A", "Di11A", "Di12A"
    }
}

def find_code_info(title):
    """
    Checks if any defined code is in the title.
    Returns (code_id, category) or attempts to extract a code-like pattern.
    If no known code found, returns a derived code from the title.
    """
    if not title:
        return "Unknown", "Custom"
    
    # First, try to match known codes
    for category, codes in CODE_DEFINITIONS.items():
        for code in codes:
            # Check if code is in title (case-insensitive or sensitive? usually codes are specific)
            if code in title:
                return code, category
    
    # If no known code found, try to extract a code-like pattern from the title
    # Common patterns: alphanumeric codes like "BGN-01", "Test-A1", etc.
    code_patterns = [
        r'\b([A-Z]{2,4}[-_]?\d{1,3}[A-Z]?)\b',  # e.g., BGN-01, TEST01A
        r'\b(\d{1,3}[-_]?[A-Z]{2,4})\b',  # e.g., 01-BGN
        r'\b([A-Z]+[-_][A-Z]+[-_]?\d*)\b',  # e.g., BGN-Reverb-1
    ]
    
    for pattern in code_patterns:
        match = re.search(pattern, title, re.IGNORECASE)
        if match:
            return match.group(1), "Custom"
    
    # If still no code found, use a shortened version of the title as the code
    # Take first 20 chars or up to first major separator
    short_title = title.split(' - ')[0] if ' - ' in title else title
    short_title = short_title.split(': ')[0] if ': ' in short_title else short_title
    short_title = short_title[:30].strip()
    
    return short_title if short_title else "Unknown", "Custom"

def format_duration(total_seconds):
    """Format duration in seconds to human readable string."""
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours}h {minutes}m {seconds}s"

def calculate_category_test_times(data, parse_date_func):
    """
    Calculate test time duration for each category.
    Handles overnight testing by calculating time per day.
    Returns a dict with category as key and time info as value.
    """
    category_times = {}
    
    for row in data:
        code_id = row[0]
        category = row[1]  # Short category (AR, RR, Di, Op)
        date_str = row[4]  # SmdDate
        
        if not category or not date_str:
            continue
            
        parsed_date = parse_date_func(date_str)
        if parsed_date == datetime.min:
            continue
        
        date_only = parsed_date.date()
        
        if category not in category_times:
            category_times[category] = {
                'min_time': parsed_date,
                'max_time': parsed_date,
                'count': 1,
                'daily_times': {date_only: {'min': parsed_date, 'max': parsed_date, 'count': 1}}
            }
        else:
            if parsed_date < category_times[category]['min_time']:
                category_times[category]['min_time'] = parsed_date
            if parsed_date > category_times[category]['max_time']:
                category_times[category]['max_time'] = parsed_date
            category_times[category]['count'] += 1
            
            # Track daily times
            if date_only not in category_times[category]['daily_times']:
                category_times[category]['daily_times'][date_only] = {'min': parsed_date, 'max': parsed_date, 'count': 1}
            else:
                if parsed_date < category_times[category]['daily_times'][date_only]['min']:
                    category_times[category]['daily_times'][date_only]['min'] = parsed_date
                if parsed_date > category_times[category]['daily_times'][date_only]['max']:
                    category_times[category]['daily_times'][date_only]['max'] = parsed_date
                category_times[category]['daily_times'][date_only]['count'] += 1
    
    # Calculate duration for each category
    result = {}
    for category, times in category_times.items():
        duration = times['max_time'] - times['min_time']
        total_seconds = int(duration.total_seconds())
        
        # Check if tests span multiple days
        daily_times = times['daily_times']
        is_overnight = len(daily_times) > 1
        
        daily_breakdown = []
        actual_total_seconds = 0
        
        if is_overnight:
            # Calculate time used each day
            for date_key in sorted(daily_times.keys()):
                day_info = daily_times[date_key]
                day_duration = day_info['max'] - day_info['min']
                day_seconds = int(day_duration.total_seconds())
                actual_total_seconds += day_seconds
                
                daily_breakdown.append({
                    'date': date_key.strftime("%m/%d/%Y"),
                    'start_time': day_info['min'].strftime("%I:%M %p"),
                    'end_time': day_info['max'].strftime("%I:%M %p"),
                    'duration': format_duration(day_seconds),
                    'duration_seconds': day_seconds,
                    'test_count': day_info['count']
                })
        else:
            actual_total_seconds = total_seconds
        
        result[category] = {
            'start_time': times['min_time'].strftime("%m/%d/%Y %I:%M %p"),
            'end_time': times['max_time'].strftime("%m/%d/%Y %I:%M %p"),
            'duration': format_duration(actual_total_seconds if is_overnight else total_seconds),
            'duration_seconds': actual_total_seconds if is_overnight else total_seconds,
            'test_count': times['count'],
            'is_overnight': is_overnight,
            'daily_breakdown': daily_breakdown
        }
    
    return result

def calculate_overall_test_time(data, parse_date_func):
    """
    Calculate overall test time across all entries regardless of category.
    Groups by date and calculates time from earliest to latest test on each day.
    Returns a dict with overall time info and daily breakdown.
    """
    if not data:
        return None
    
    daily_times = {}
    total_count = 0
    
    for row in data:
        date_str = row[4]  # SmdDate
        
        if not date_str:
            continue
            
        parsed_date = parse_date_func(date_str)
        if parsed_date == datetime.min:
            continue
        
        date_only = parsed_date.date()
        total_count += 1
        
        if date_only not in daily_times:
            daily_times[date_only] = {'min': parsed_date, 'max': parsed_date, 'count': 1}
        else:
            if parsed_date < daily_times[date_only]['min']:
                daily_times[date_only]['min'] = parsed_date
            if parsed_date > daily_times[date_only]['max']:
                daily_times[date_only]['max'] = parsed_date
            daily_times[date_only]['count'] += 1
    
    if not daily_times:
        return None
    
    # Calculate time used each day and sum up
    daily_breakdown = []
    total_seconds = 0
    
    for date_key in sorted(daily_times.keys()):
        day_info = daily_times[date_key]
        day_duration = day_info['max'] - day_info['min']
        day_seconds = int(day_duration.total_seconds())
        total_seconds += day_seconds
        
        daily_breakdown.append({
            'date': date_key.strftime("%m/%d/%Y"),
            'start_time': day_info['min'].strftime("%I:%M %p"),
            'end_time': day_info['max'].strftime("%I:%M %p"),
            'duration': format_duration(day_seconds),
            'duration_seconds': day_seconds,
            'test_count': day_info['count']
        })
    
    # Get overall min and max times
    all_min = min(daily_times[d]['min'] for d in daily_times)
    all_max = max(daily_times[d]['max'] for d in daily_times)
    
    return {
        'start_time': all_min.strftime("%m/%d/%Y %I:%M %p"),
        'end_time': all_max.strftime("%m/%d/%Y %I:%M %p"),
        'duration': format_duration(total_seconds),
        'duration_seconds': total_seconds,
        'test_count': total_count,
        'num_days': len(daily_times),
        'daily_breakdown': daily_breakdown
    }

def extract_54db_noise_results(file_paths):
    """
    Extract 54dB noise scenario SMOS, NMOS, GMOS results from Status Overview table.
    Looks for rows with 54dB patterns and extracts MOS values.
    Returns list of dicts with device, NS setting, and MOS values.
    """
    results = []
    
    for file_path in file_paths:
        try:
            doc = Document(file_path)
            file_name = os.path.basename(file_path)
            device_name = file_name.replace('.docx', '')
            
            # Extract report time and lab name from document
            report_time = ""
            lab_name = ""
            
            # Check headers and footers for lab name (AST or PAL)
            for section in doc.sections:
                # Check header
                if section.header:
                    header_text = ""
                    for para in section.header.paragraphs:
                        header_text += para.text + " "
                    header_lower = header_text.lower()
                    if "ast" in header_lower:
                        lab_name = "AST"
                    elif "pal" in header_lower:
                        lab_name = "PAL"
                
                # Check footer if lab not found in header
                if not lab_name and section.footer:
                    footer_text = ""
                    for para in section.footer.paragraphs:
                        footer_text += para.text + " "
                    footer_lower = footer_text.lower()
                    if "ast" in footer_lower:
                        lab_name = "AST"
                    elif "pal" in footer_lower:
                        lab_name = "PAL"
                
                if lab_name:
                    break
            
            # Check paragraphs for lab name and report time (date only)
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not report_time:
                    # Extract date only (no time)
                    date_match = re.search(r"(\d{1,2}/\d{1,2}/\d{4})", text)
                    if date_match:
                        report_time = date_match.group(1)
                # Also check paragraphs for lab name if not found in header/footer
                if not lab_name:
                    text_lower = text.lower()
                    if "ast" in text_lower:
                        lab_name = "AST"
                    elif "pal" in text_lower:
                        lab_name = "PAL"
            
            # Track 54dB results for this file - group by NS setting
            file_results = {}  # key: NS_Setting, value: dict with MOS values
            
            # Process tables to find the Status Overview table
            for table in doc.tables:
                header_row = []
                is_status_table = False
                
                # Check first row for expected headers
                if len(table.rows) > 0:
                    first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    first_row_text = " ".join(first_row_cells).lower()
                    
                    # Check if this is the status table
                    if ("smd" in first_row_text and "status" in first_row_text and 
                        ("single value" in first_row_text or "description" in first_row_text)):
                        is_status_table = True
                        header_row = first_row_cells
                
                if not is_status_table:
                    continue
                
                # Find column indices
                unique_headers = []
                prev_text = None
                for idx, header in enumerate(header_row):
                    if header != prev_text:
                        unique_headers.append((idx, header))
                    prev_text = header
                
                smd_idx = status_idx = desc_idx = value_idx = -1
                
                for idx, header in unique_headers:
                    header_normalized = ' '.join(header.split()).lower().strip()
                    header_lower = header.lower().strip()
                    
                    if "smd" in header_lower and smd_idx == -1:
                        smd_idx = idx
                    elif header_normalized == "status" and status_idx == -1:
                        status_idx = idx
                    elif "description" in header_lower and desc_idx == -1:
                        desc_idx = idx
                    elif ("single value" in header_normalized or "single\nvalue" in header_lower) and "description" not in header_lower and value_idx == -1:
                        value_idx = idx
                
                # Fallback for value column
                if value_idx == -1 and len(unique_headers) >= 4:
                    last_header = unique_headers[-1][1].lower().strip()
                    if "object" in last_header and len(unique_headers) >= 5:
                        value_idx = unique_headers[-2][0]
                    else:
                        value_idx = unique_headers[3][0] if len(unique_headers) > 3 else unique_headers[-1][0]
                
                # Process data rows
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:  # Skip header row
                        continue
                    
                    row_cells = [cell.text.strip() for cell in row.cells]
                    
                    # Handle merged cells
                    unique_cells = []
                    prev_text = None
                    for cell_text in row_cells:
                        if cell_text != prev_text:
                            unique_cells.append(cell_text)
                        prev_text = cell_text
                    
                    # Extract values
                    smd_val = row_cells[smd_idx] if smd_idx >= 0 and smd_idx < len(row_cells) else ""
                    desc_val = row_cells[desc_idx] if desc_idx >= 0 and desc_idx < len(row_cells) else ""
                    single_val = row_cells[value_idx] if value_idx >= 0 and value_idx < len(row_cells) else ""
                    
                    if not single_val and len(unique_cells) >= 4:
                        single_val = unique_cells[3] if unique_cells[3] != desc_val else ""
                    
                    if not smd_val:
                        continue
                    
                    smd_lower = smd_val.lower()
                    desc_lower = desc_val.lower()
                    
                    # Check if this is a 54dB noise scenario row
                    patterns_54db = ["54db", "54 db", "54dba", "54 dba", "bgn_54", "bgn54", "bgn-54"]
                    is_54db = any(pattern in smd_lower.replace(" ", "") or pattern in smd_lower for pattern in patterns_54db)
                    
                    if not is_54db:
                        continue
                    
                    # Determine NS setting
                    ns_setting = "Android"  # Default
                    ns_patterns = [
                        ("ns on", "NS ON"), ("ns_on", "NS ON"), ("ns-on", "NS ON"), ("ns=on", "NS ON"),
                        ("ns off", "NS OFF"), ("ns_off", "NS OFF"), ("ns-off", "NS OFF"), ("ns=off", "NS OFF"),
                        ("nson", "NS ON"), ("nsoff", "NS OFF")
                    ]
                    smd_no_space = smd_lower.replace(" ", "")
                    for pattern, setting in ns_patterns:
                        if pattern in smd_lower or pattern in smd_no_space:
                            ns_setting = setting
                            break
                    
                    # Determine if 2ndTalker or BGN
                    # Normalize SMD for pattern matching (remove newlines, extra spaces)
                    smd_normalized = smd_lower.replace("\n", " ").replace("  ", " ").strip()
                    smd_no_space = smd_lower.replace("\n", "").replace(" ", "")
                    
                    # Check for 2ndTalker patterns
                    talker_patterns = ["art mouth", "artmouth", "mouth", "2 mouth", "2nd mouth", "2mouth", "2nd", "talker", "2"]
                    is_2nd_talker = any(pattern in smd_normalized or pattern in smd_no_space for pattern in talker_patterns)
                    
                    is_bgn = any(x in smd_lower for x in ["bgn", "hae-bgn", "haebgn", "3pass"])
                    
                    # Determine MOS type from SMD and Description
                    mos_type = None
                    if "s-mos" in smd_lower or "smos" in smd_lower:
                        mos_type = "SMOS"
                    elif "n-mos" in smd_lower or "nmos" in smd_lower:
                        mos_type = "NMOS"
                    elif "g-mos" in desc_lower or "gmos" in desc_lower:
                        mos_type = "GMOS"
                    
                    # Extract numeric MOS value
                    mos_value = ""
                    if single_val:
                        val_clean = single_val.replace(",", ".").strip()
                        match = re.search(r"([\d.]+)", val_clean)
                        if match:
                            mos_value = match.group(1)
                    
                    # Initialize result entry for this NS setting if not exists
                    if ns_setting not in file_results:
                        file_results[ns_setting] = {
                            'Device': device_name,
                            'Lab': lab_name,
                            'Report_Time': report_time,
                            'NS_Setting': ns_setting,
                            'SMOS_2ndTalker': '',
                            'NMOS_2ndTalker': '',
                            'GMOS_2ndTalker': '',
                            'SMOS_BGN': '',
                            'NMOS_BGN': '',
                            'GMOS_BGN': '',
                            'FilePath': file_path
                        }
                    
                    # Assign MOS value to appropriate field
                    if mos_type and mos_value:
                        if is_2nd_talker:
                            if mos_type == "SMOS":
                                file_results[ns_setting]['SMOS_2ndTalker'] = mos_value
                            elif mos_type == "NMOS":
                                file_results[ns_setting]['NMOS_2ndTalker'] = mos_value
                            elif mos_type == "GMOS":
                                file_results[ns_setting]['GMOS_2ndTalker'] = mos_value
                        elif is_bgn:
                            if mos_type == "SMOS":
                                file_results[ns_setting]['SMOS_BGN'] = mos_value
                            elif mos_type == "NMOS":
                                file_results[ns_setting]['NMOS_BGN'] = mos_value
                            elif mos_type == "GMOS":
                                file_results[ns_setting]['GMOS_BGN'] = mos_value
                        else:
                            # Default to BGN if pattern not clear
                            if mos_type == "SMOS":
                                file_results[ns_setting]['SMOS_BGN'] = mos_value
                            elif mos_type == "NMOS":
                                file_results[ns_setting]['NMOS_BGN'] = mos_value
                            elif mos_type == "GMOS":
                                file_results[ns_setting]['GMOS_BGN'] = mos_value
            
            # Add file results to main results list
            for ns_setting, result_entry in file_results.items():
                results.append(result_entry)
            
            # Debug output
            if file_results:
                print(f"  Found {len(file_results)} 54dB noise scenario entries in {file_name}")
            else:
                print(f"  No 54dB noise scenario found in {file_name}")
                
        except Exception as e:
            print(f"Error extracting 54dB noise results from {os.path.basename(file_path)}: {str(e)}")
    
    return results

def extract_double_talk_performance(file_paths):
    """
    Extract double talk performance data from Status Overview table.
    Looks for rows where Single Value Description contains "Attenuation during double talk [dB]".
    Returns list of dicts with device, SMD, status, description, and single value.
    """
    results = []
    
    for file_path in file_paths:
        try:
            doc = Document(file_path)
            file_name = os.path.basename(file_path)
            device_name = file_name.replace('.docx', '')
            
            # Process tables to find the Status Overview table
            for table in doc.tables:
                header_row = []
                is_status_table = False
                
                # Check first row for expected headers
                if len(table.rows) > 0:
                    first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    first_row_text = " ".join(first_row_cells).lower()
                    
                    # Check if this is the status table
                    if ("smd" in first_row_text and "status" in first_row_text and 
                        ("single value" in first_row_text or "description" in first_row_text)):
                        is_status_table = True
                        header_row = first_row_cells
                
                if not is_status_table:
                    continue
                
                # Find column indices - handle merged cells
                unique_headers = []
                prev_text = None
                for idx, header in enumerate(header_row):
                    if header != prev_text:
                        unique_headers.append((idx, header))
                    prev_text = header
                
                smd_idx = status_idx = desc_idx = value_idx = -1
                
                for idx, header in unique_headers:
                    header_normalized = ' '.join(header.split()).lower().strip()
                    header_lower = header.lower().strip()
                    
                    if "smd" in header_lower and smd_idx == -1:
                        smd_idx = idx
                    elif header_normalized == "status" and status_idx == -1:
                        status_idx = idx
                    elif "description" in header_lower and desc_idx == -1:
                        desc_idx = idx
                    elif ("single value" in header_normalized or "single\nvalue" in header_lower) and "description" not in header_lower and value_idx == -1:
                        value_idx = idx
                
                # Fallback for value column
                if value_idx == -1 and len(unique_headers) >= 4:
                    last_header = unique_headers[-1][1].lower().strip()
                    if "object" in last_header and len(unique_headers) >= 5:
                        value_idx = unique_headers[-2][0]
                    else:
                        value_idx = unique_headers[3][0] if len(unique_headers) > 3 else unique_headers[-1][0]
                
                # Process data rows
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:  # Skip header row
                        continue
                    
                    row_cells = [cell.text.strip() for cell in row.cells]
                    
                    # Handle merged cells
                    unique_cells = []
                    prev_text = None
                    for cell_text in row_cells:
                        if cell_text != prev_text:
                            unique_cells.append(cell_text)
                        prev_text = cell_text
                    
                    # Extract values
                    smd_val = row_cells[smd_idx] if smd_idx >= 0 and smd_idx < len(row_cells) else ""
                    status_val = row_cells[status_idx] if status_idx >= 0 and status_idx < len(row_cells) else ""
                    desc_val = row_cells[desc_idx] if desc_idx >= 0 and desc_idx < len(row_cells) else ""
                    single_val = row_cells[value_idx] if value_idx >= 0 and value_idx < len(row_cells) else ""
                    
                    if not single_val and len(unique_cells) >= 4:
                        single_val = unique_cells[3] if unique_cells[3] != desc_val else ""
                    
                    # Check if this is a double talk attenuation row
                    if "attenuation during double talk" in desc_val.lower():
                        results.append({
                            'FileName': file_name,
                            'Device': device_name,
                            'SMD': smd_val,
                            'Status': status_val,
                            'Description': desc_val,
                            'SingleValue': single_val,
                            'FilePath': file_path
                        })
                        
        except Exception as e:
            print(f"Error extracting double talk data from {os.path.basename(file_path)}: {str(e)}")
    
    return results

def extract_smd_settings(file_paths):
    """
    Extract equipment settings from SmdSettings style paragraphs.
    Returns dict with labCORE info, HATS info, BEQ settings, etc.
    
    Expected format (tab-separated):
    - labCORE serial	77000079	Nickname	Reverb1 LC
    - Firmware	3.11.9
    - BEQ Settings: Equalization	DF, HATS serial	12309013, Pinna	Type 3.3
    - Artificial Head: Ser. Nr.	12309013	Pinna type	Type 3.3
    """
    settings_data = {
        'labCORE': [],  # List of {serial, firmware, nickname}
        'HATS': [],     # List of {serial, pinna_type}
        'BEQ': [],      # List of {equalization, test_code}
    }
    
    for file_path in file_paths:
        try:
            doc = Document(file_path)
            file_name = os.path.basename(file_path)
            
            current_labcore = {}
            current_hats = {}
            current_beq = {}
            current_test_code = ""
            in_labcore_section = False
            in_beq_section = False
            in_hats_section = False
            
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                text_lower = text.lower()
                
                # Track current test code from SmdTitle
                if style_name == "SmdTitle":
                    # Extract test code like P05R, P10R
                    code_match = re.search(r'\b(P\d{2}[AR]|Di\d{2}[A]|Op[A-Z]\d{2}[AR]?)\b', text)
                    if code_match:
                        current_test_code = code_match.group(1)
                
                # Check for SmdSetting style OR content patterns (more flexible matching)
                is_settings_paragraph = (
                    style_name == "SmdSetting" or
                    "labcore settings" in text_lower or
                    "labcore serial" in text_lower or
                    "beq settings" in text_lower or
                    "artificial head" in text_lower or
                    "hats serial" in text_lower or
                    "equalization" in text_lower or
                    "pinna type" in text_lower or
                    "ser. nr." in text_lower or
                    ("firmware" in text_lower and "sync" in text_lower)
                )
                
                if is_settings_paragraph:
                    # Detect section headers
                    if "labcore settings" in text_lower:
                        in_labcore_section = True
                        in_beq_section = False
                        in_hats_section = False
                        continue
                    elif "beq settings" in text_lower:
                        in_beq_section = True
                        in_labcore_section = False
                        in_hats_section = False
                        continue
                    elif "artificial head" in text_lower or ("hats" in text_lower and "settings" in text_lower):
                        in_hats_section = True
                        in_labcore_section = False
                        in_beq_section = False
                        continue
                    elif text.startswith("----------"):
                        # Section separator - reset sections
                        in_labcore_section = False
                        in_beq_section = False
                        in_hats_section = False
                        continue
                    
                    # Parse labCORE section content
                    # Format: "labCORE serial	77000079	Nickname	Reverb1 LC"
                    # Format: "Firmware	3.11.9	Sync source	Internal"
                    if in_labcore_section or "labcore serial" in text_lower or ("firmware" in text_lower and "sync" in text_lower):
                        # Look for tab-separated values
                        # Pattern: labCORE serial<tab>VALUE<tab>Nickname<tab>VALUE
                        serial_match = re.search(r'labcore\s+serial[\t\s]+(\d+)', text, re.IGNORECASE)
                        if serial_match:
                            current_labcore['serial'] = serial_match.group(1).strip()
                        
                        # Nickname in same line
                        nick_match = re.search(r'nickname[\t\s]+([^\t\n]+)', text, re.IGNORECASE)
                        if nick_match:
                            current_labcore['nickname'] = nick_match.group(1).strip()
                        
                        # Firmware (might be on separate line)
                        fw_match = re.search(r'firmware[\t\s]+([0-9.]+)', text, re.IGNORECASE)
                        if fw_match:
                            current_labcore['firmware'] = fw_match.group(1).strip()
                    
                    # Parse BEQ section content
                    # Format: "Block mode	Active	Equalization	DF"
                    # Format: "HATS serial	12309013	Pinna	Type 3.3"
                    if in_beq_section or "equalization" in text_lower or "hats serial" in text_lower:
                        # Equalization value
                        eq_match = re.search(r'equalization[\t\s]+([^\t\n]+)', text, re.IGNORECASE)
                        if eq_match:
                            eq_value = eq_match.group(1).strip()
                            current_beq['equalization'] = eq_value
                            current_beq['test_code'] = current_test_code
                            # Also store equalization with HATS info
                            current_hats['equalization'] = eq_value
                            # Check if DF is mentioned
                            if 'df' in eq_value.lower() or 'diffuse' in eq_value.lower():
                                current_beq['has_df'] = True
                        
                        # HATS serial in BEQ section
                        hats_serial_match = re.search(r'hats\s+serial[\t\s]+(\d+)', text, re.IGNORECASE)
                        if hats_serial_match:
                            current_beq['hats_serial'] = hats_serial_match.group(1).strip()
                            current_hats['serial'] = hats_serial_match.group(1).strip()
                        
                        # Pinna in BEQ section (just "Pinna" not "Pinna type")
                        pinna_match = re.search(r'\bpinna[\t\s]+([^\t\n]+)', text, re.IGNORECASE)
                        if pinna_match and "pinna type" not in text_lower:
                            current_beq['pinna'] = pinna_match.group(1).strip()
                            current_hats['pinna'] = pinna_match.group(1).strip()
                    
                    # Parse Artificial Head section
                    # Format: "Ser. Nr.	12309013	Pinna type	Type 3.3"
                    if in_hats_section or "ser. nr." in text_lower or "pinna type" in text_lower:
                        # Serial number
                        ser_match = re.search(r'ser\.?\s*nr\.?[\t\s]+(\d+)', text, re.IGNORECASE)
                        if ser_match:
                            current_hats['serial'] = ser_match.group(1).strip()
                        
                        # Pinna type
                        pinna_match = re.search(r'pinna\s*type[\t\s]+([^\t\n]+)', text, re.IGNORECASE)
                        if pinna_match:
                            current_hats['pinna'] = pinna_match.group(1).strip()
            
            # Add collected data
            if current_labcore:
                current_labcore['file'] = file_name
                settings_data['labCORE'].append(current_labcore)
            
            if current_hats:
                current_hats['file'] = file_name
                settings_data['HATS'].append(current_hats)
            
            if current_beq:
                current_beq['file'] = file_name
                settings_data['BEQ'].append(current_beq)
            
            # Debug output
            if current_labcore or current_hats or current_beq:
                print(f"  Found SmdSettings: labCORE={bool(current_labcore)}, HATS={bool(current_hats)}, BEQ={bool(current_beq)}")
                
        except Exception as e:
            print(f"Error extracting SmdSettings from {os.path.basename(file_path)}: {str(e)}")
    
    return settings_data

def extract_status_table(file_paths):
    """
    Extract Status Overview table content from Word documents.
    Looks for tables with header "SMD	Status	Single Value Description	Single Value"
    Also extracts limit values from Limits tables (with SmdLimitsTableHeader, SmdLimitsTableData styles).
    Returns all rows and identifies "Not OK" status entries with limits.
    """
    all_status_rows = []
    not_ok_rows = []
    
    for file_path in file_paths:
        try:
            doc = Document(file_path)
            file_name = os.path.basename(file_path)
            
            # Build a map of SMD title -> limits by scanning document structure
            smd_limits_map = {}
            current_smd_title = None
            
            # Get document body elements in order to track SmdTitle -> Limits table association
            body_elements = doc.element.body
            
            for element in body_elements:
                # Check if this is a paragraph
                if element.tag.endswith('p'):
                    for para in doc.paragraphs:
                        if para._element is element:
                            if para.style.name == "SmdTitle":
                                current_smd_title = para.text.strip()
                            break
                
                # Check if this is a table - look for limit styles in table cells
                elif element.tag.endswith('tbl'):
                    if current_smd_title:
                        for table in doc.tables:
                            if table._element is element:
                                # Check table cells for SmdLimits* styles
                                limit_type = ""
                                limit_value = ""
                                is_limits_table = False
                                
                                for row in table.rows:
                                    for cell in row.cells:
                                        for para in cell.paragraphs:
                                            style_name = para.style.name
                                            text = para.text.strip()
                                            
                                            # Check for SmdLimitsHeader or header styles
                                            if style_name in ["SmdLimitsHeader", "SmdLimitsTableHeader", "SmdLimtsTableHeader"]:
                                                is_limits_table = True
                                                text_lower = text.lower()
                                                if "upper" in text_lower:
                                                    limit_type = "Upper"
                                                elif "lower" in text_lower:
                                                    limit_type = "Lower"
                                            
                                            # Check for data/text styles with limit values
                                            elif style_name in ["SmdLimitsTableData", "SmdLimitsTableText", "SmdLimitsText"]:
                                                is_limits_table = True
                                                # Extract numeric value with optional unit
                                                val_match = re.search(r"([+-]?[\d.]+)\s*(dB|db|DB|%|ms|Hz|s)?", text)
                                                if val_match and not text.lower().startswith("run"):
                                                    if val_match.group(2):
                                                        limit_value = f"{val_match.group(1)} {val_match.group(2)}"
                                                    else:
                                                        limit_value = val_match.group(1)
                                
                                # Also check by content if styles weren't found
                                if not is_limits_table and len(table.rows) >= 2:
                                    first_row_text = " ".join([cell.text.strip().lower() for cell in table.rows[0].cells])
                                    if "limits" in first_row_text:
                                        is_limits_table = True
                                        if "upper" in first_row_text:
                                            limit_type = "Upper"
                                        elif "lower" in first_row_text:
                                            limit_type = "Lower"
                                        
                                        # Get value from data rows
                                        for row in table.rows[1:]:
                                            for cell in row.cells:
                                                text = cell.text.strip()
                                                if not text.lower().startswith("run"):
                                                    val_match = re.search(r"([+-]?[\d.]+)\s*(dB|db|DB|%|ms|Hz|s)?", text)
                                                    if val_match:
                                                        if val_match.group(2):
                                                            limit_value = f"{val_match.group(1)} {val_match.group(2)}"
                                                        else:
                                                            limit_value = val_match.group(1)
                                                        break
                                            if limit_value:
                                                break
                                
                                if is_limits_table and limit_type and limit_value:
                                    if current_smd_title not in smd_limits_map:
                                        smd_limits_map[current_smd_title] = []
                                    limit_info = f"{limit_type}: {limit_value}"
                                    if limit_info not in smd_limits_map[current_smd_title]:
                                        smd_limits_map[current_smd_title].append(limit_info)
                                break
            
            # Debug output
            if smd_limits_map:
                print(f"  Found limits for {len(smd_limits_map)} test cases")
            else:
                print(f"  No limits found in document")
            
            # Process tables to find the status table
            for table in doc.tables:
                header_row = []
                is_status_table = False
                
                # Check first row for expected headers
                if len(table.rows) > 0:
                    first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    first_row_text = " ".join(first_row_cells).lower()
                    
                    # Check if this is the status table
                    if ("smd" in first_row_text and "status" in first_row_text and 
                        ("single value" in first_row_text or "description" in first_row_text)):
                        is_status_table = True
                        header_row = first_row_cells
                
                if not is_status_table:
                    continue
                
                # Debug: print header info
                print(f"  Status table found with {len(header_row)} columns: {header_row}")
                
                # Handle merged cells - get unique headers and their actual column positions
                # In Word tables with merged cells, cell.text may repeat for merged areas
                unique_headers = []
                header_positions = {}
                prev_text = None
                for idx, header in enumerate(header_row):
                    if header != prev_text:  # New unique header
                        unique_headers.append((idx, header))
                        header_positions[header.lower().strip()] = idx
                    prev_text = header
                
                # Find column indices - improved detection with merged cell handling
                # Handle multi-line headers like "Single\nValue"
                smd_idx = status_idx = desc_idx = value_idx = -1
                
                for idx, header in unique_headers:
                    # Normalize header: replace newlines with space, collapse whitespace
                    header_normalized = ' '.join(header.split()).lower().strip()
                    header_lower = header.lower().strip()
                    
                    if "smd" in header_lower and smd_idx == -1:
                        smd_idx = idx
                    elif header_normalized == "status" and status_idx == -1:
                        status_idx = idx
                    elif "description" in header_lower and desc_idx == -1:
                        desc_idx = idx
                    elif ("single value" in header_normalized or "single\nvalue" in header_lower) and "description" not in header_lower and value_idx == -1:
                        value_idx = idx
                
                # If "Single Value" column not found, try finding by position (usually 4th column, before optional "Object")
                if value_idx == -1 and len(unique_headers) >= 4:
                    # Check if last column is "Object" - if so, Single Value is second to last
                    last_header = unique_headers[-1][1].lower().strip()
                    if "object" in last_header and len(unique_headers) >= 5:
                        value_idx = unique_headers[-2][0]  # Second to last column
                        print(f"  Single Value column not explicitly found, using second-to-last (before Object): {value_idx}")
                    else:
                        # The 4th unique column is typically Single Value
                        value_idx = unique_headers[3][0] if len(unique_headers) > 3 else unique_headers[-1][0]
                        print(f"  Single Value column not explicitly found, using column index: {value_idx}")
                
                # Debug: print column indices
                print(f"  Column indices: SMD={smd_idx}, Status={status_idx}, Desc={desc_idx}, Value={value_idx}")
                
                # Helper function to normalize text for comparison
                def normalize_for_match(text):
                    """Remove Index info and extra suffixes for matching."""
                    if not text:
                        return ""
                    # Remove ", Index: X" or ",Index: X" patterns
                    normalized = re.sub(r',?\s*Index:\s*\d+', '', text, flags=re.IGNORECASE)
                    # Remove trailing whitespace
                    normalized = normalized.strip()
                    return normalized
                
                # Helper function for similarity matching
                def get_match_score(smd_val, smd_title):
                    """Calculate match score between SMD value and SmdTitle."""
                    # Normalize both strings
                    norm_smd = normalize_for_match(smd_val).lower()
                    norm_title = normalize_for_match(smd_title).lower()
                    
                    if not norm_smd or not norm_title:
                        return 0
                    
                    # Exact match after normalization
                    if norm_smd == norm_title:
                        return 100
                    
                    # One contains the other
                    if norm_smd in norm_title:
                        return 90
                    if norm_title in norm_smd:
                        return 85
                    
                    # Extract test code (e.g., P02A, P07A) and compare
                    code_pattern = r'^([A-Z]{1,3}\d{1,3}[A-Z]?)'
                    smd_code_match = re.match(code_pattern, norm_smd)
                    title_code_match = re.match(code_pattern, norm_title)
                    
                    if smd_code_match and title_code_match:
                        if smd_code_match.group(1) == title_code_match.group(1):
                            # Same test code, check rest of text
                            smd_rest = norm_smd[len(smd_code_match.group(1)):].strip()
                            title_rest = norm_title[len(title_code_match.group(1)):].strip()
                            
                            # Remove common suffixes like "NS ON", "NS OFF", etc.
                            suffixes_to_remove = [' ns on', ' ns off', ' ns_on', ' ns_off']
                            for suffix in suffixes_to_remove:
                                smd_rest = smd_rest.replace(suffix, '')
                                title_rest = title_rest.replace(suffix, '')
                            
                            smd_rest = smd_rest.strip()
                            title_rest = title_rest.strip()
                            
                            if smd_rest == title_rest:
                                return 95
                            if smd_rest in title_rest or title_rest in smd_rest:
                                return 80
                            
                            # Check word overlap
                            smd_words = set(smd_rest.split())
                            title_words = set(title_rest.split())
                            if smd_words and title_words:
                                overlap = len(smd_words & title_words)
                                total = len(smd_words | title_words)
                                if total > 0:
                                    return int(70 * overlap / total)
                    
                    return 0
                
                # Process data rows
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:  # Skip header row
                        continue
                    
                    row_cells = [cell.text.strip() for cell in row.cells]
                    
                    # Handle merged cells in data rows - get unique cell values
                    unique_cells = []
                    prev_text = None
                    for cell_text in row_cells:
                        if cell_text != prev_text:  # Only add if different from previous
                            unique_cells.append(cell_text)
                        prev_text = cell_text
                    
                    # Extract values based on column indices
                    # Try using unique_cells if row_cells at value_idx seems to be duplicated
                    smd_val = row_cells[smd_idx] if smd_idx >= 0 and smd_idx < len(row_cells) else ""
                    status_val = row_cells[status_idx] if status_idx >= 0 and status_idx < len(row_cells) else ""
                    desc_val = row_cells[desc_idx] if desc_idx >= 0 and desc_idx < len(row_cells) else ""
                    single_val = row_cells[value_idx] if value_idx >= 0 and value_idx < len(row_cells) else ""
                    
                    # If single_val is empty but we have 4 unique cells, try the 4th unique cell
                    if not single_val and len(unique_cells) >= 4:
                        single_val = unique_cells[3] if unique_cells[3] != desc_val else ""
                    
                    # Debug: print first few Not OK rows to see structure
                    if status_val.lower() == "not ok" and row_idx <= 3:
                        print(f"    Row {row_idx}: cells={row_cells}")
                        print(f"    Row {row_idx}: unique_cells={unique_cells}")
                        print(f"    Row {row_idx}: smd={smd_val}, status={status_val}, desc={desc_val}, single_val={single_val}")
                    
                    # Skip empty rows
                    if not smd_val and not status_val:
                        continue
                    
                    # Find matching limits for this SMD using similarity matching
                    limit_val = ""
                    best_score = 0
                    best_match_title = None
                    
                    for smd_title, limits in smd_limits_map.items():
                        score = get_match_score(smd_val, smd_title)
                        if score > best_score:
                            best_score = score
                            best_match_title = smd_title
                    
                    # Use match if score is good enough (threshold: 70)
                    if best_score >= 70 and best_match_title:
                        limit_val = " | ".join(smd_limits_map[best_match_title])
                    
                    row_entry = {
                        'FileName': file_name,
                        'SMD': smd_val,
                        'Status': status_val,
                        'SingleValueDescription': desc_val,
                        'SingleValue': single_val,
                        'Limit': limit_val,
                        'FilePath': file_path
                    }
                    
                    all_status_rows.append(row_entry)
                    
                    # Check for "Not OK" status
                    if status_val.lower() == "not ok" or "not ok" in status_val.lower():
                        not_ok_rows.append(row_entry)
                
        except Exception as e:
            print(f"Error extracting status table from {os.path.basename(file_path)}: {str(e)}")
    
    return all_status_rows, not_ok_rows

def process_reports(file_paths):
    extracted_rows = []
    is_shared_speakerphone = False

    print(f"Processing {len(file_paths)} file(s)...")

    for file_path in file_paths:
        print(f"Reading file: {file_path}")
        try:
            doc = Document(file_path)
            
            # Temporary lists to hold data for the current file
            titles = []
            dates = []

            # Track styles found for debugging
            styles_found = set()
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                # Check for "Shared" or "Speakerphone" keyword to identify test case requirements
                if not is_shared_speakerphone:
                    if "Shared" in text or "Speakerphone" in text:
                        is_shared_speakerphone = True

                style_name = paragraph.style.name
                styles_found.add(style_name)

                if style_name == "SmdTitle":
                    titles.append(text)
                    continue

                # Check if this line is the "Unmodified HEAD acoustics Measurement Descriptor" marker
                # The user indicated this line might effectively label the PREVIOUS line as the date,
                # and this line itself might wrongly have SmdDate style.
                if "Unmodified HEAD acoustics Measurement Descriptor" in text:
                    if i > 0:
                        prev_para = doc.paragraphs[i-1]
                        # If previous line was SmdDate, it was already processed in previous iteration.
                        # If NOT SmdDate, we must extract date from it now.
                        if prev_para.style.name != "SmdDate":
                            clean_date = extract_clean_date(prev_para.text)
                            if clean_date:
                                dates.append(clean_date)
                    # Skip this paragraph so we don't treat the marker text as a date
                    continue

                if style_name == "SmdDate":
                    clean_date = extract_clean_date(text)
                    if clean_date:
                        dates.append(clean_date)
            
            # Debug output: show what was found
            print(f"  Found {len(titles)} SmdTitle entries, {len(dates)} SmdDate entries")
            if titles:
                print(f"  Sample titles: {titles[:3]}")
            if not titles and not dates:
                print(f"  Styles found in document: {sorted(styles_found)}")
            
            # Align titles and dates found in this file
            # Assuming data appears in pairs or lists. We will match by index.
            max_items = max(len(titles), len(dates)) if titles or dates else 0
            
            # Track how many have known vs custom codes
            known_codes_count = 0
            custom_codes_count = 0
            
            for i in range(max_items):
                title_text = titles[i] if i < len(titles) else ""
                date_text = dates[i] if i < len(dates) else ""
                
                # Check for CodeID in SmdTitle (now always returns something)
                code_id, category = find_code_info(title_text)
                
                if category == "Custom":
                    custom_codes_count += 1
                else:
                    known_codes_count += 1
                
                # Store as dict for post-processing (now includes all entries)
                extracted_rows.append({
                    "CodeID": code_id,
                    "Category": category,
                    "SmdTitle": title_text,
                    "SmdDate": date_text,
                    "FilePath": file_path
                })
            
            if max_items > 0:
                print(f"  Added {max_items} entries ({known_codes_count} known codes, {custom_codes_count} custom)")
                    
        except Exception as e:
            print(f"Failed to process file: {os.path.basename(file_path)}")
            print(f"Error: {str(e)}")

    # Post-process: Sort by SmdTitle then SmdDate, and add SmdTitle_Times
    print("Sorting and indexing data...")
    
    # helper for parsing date
    def parse_date(date_str):
        try:
            return datetime.strptime(date_str, "%m/%d/%Y %I:%M %p")
        except (ValueError, TypeError):
            return datetime.min

    # Sort primarily by SmdTitle, secondarily by Date
    extracted_rows.sort(key=lambda x: (x["SmdTitle"], parse_date(x["SmdDate"])))

    # Assign SmdTitle_Times
    final_output = []
    title_counts = {}
    
    category_map = {
        "P-series (A)": "AR",
        "P-series (R)": "RR",
        "Device-direct (Di)": "Di",
        "Option codes": "Op",
        "Custom": "Custom"
    }

    for row in extracted_rows:
        title = row["SmdTitle"]
        if title not in title_counts:
            title_counts[title] = 0
        title_counts[title] += 1
        
        short_cat = category_map.get(row["Category"], "")
        
        final_output.append([
            row["CodeID"],
            short_cat,
            title_counts[title], # SmdTitle_Times
            row["SmdTitle"],
            row["SmdDate"],
            row["FilePath"]
        ])
    
    # Extract 54dB noise scenario results
    noise_54db_results = extract_54db_noise_results(file_paths)
    
    # Extract status table results
    all_status_rows, not_ok_rows = extract_status_table(file_paths)
    
    # Extract double talk performance data
    double_talk_results = extract_double_talk_performance(file_paths)
    
    # Extract SmdSettings (equipment info)
    smd_settings = extract_smd_settings(file_paths)
    
    # Extract ACQUA and Database version info
    acqua_db_info = extract_acqua_database_info(file_paths)

    return final_output, is_shared_speakerphone, noise_54db_results, all_status_rows, not_ok_rows, double_talk_results, smd_settings, acqua_db_info

def main():
    # Display version info
    print(get_version_info())
    
    # Set up the root window and hide it (we only want the file dialog)
    root = tk.Tk()
    root.withdraw()

    print("Please select the ACQUA audio report Word files (.docx)...")
    
    # Open file selector
    file_paths = filedialog.askopenfilenames(
        title="Select ACQUA Audio Reports",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
    )

    if not file_paths:
        print("No files selected. Exiting.")
        return

    # Process files
    data, is_shared_speakerphone, noise_54db_results, all_status_rows, not_ok_rows, double_talk_results, smd_settings, acqua_db_info = process_reports(file_paths)

    if not data:
        print("No matching 'SmdTitle' or 'SmdDate' styles found in the selected files.")
        return
    
    # Helper function for parsing date (same as in process_reports)
    def parse_date(date_str):
        try:
            return datetime.strptime(date_str, "%m/%d/%Y %I:%M %p")
        except (ValueError, TypeError):
            return datetime.min
    
    # Calculate and display test times by category
    category_times = calculate_category_test_times(data, parse_date)
    
    test_time_output = []
    if category_times:
        print(f"\n{'='*60}")
        header_msg = "--- Test Time by Category ---"
        print(f"{header_msg}")
        print(f"{'='*60}")
        test_time_output.append([header_msg])
        test_time_output.append(["Category", "Test Count", "Start Time", "End Time", "Duration", "Daily Breakdown"])
        
        # Print formatted table header
        print(f"\n{'Category':<20} {'Tests':<8} {'Duration':<15} {'Time Range':<30}")
        print(f"{'-'*20} {'-'*8} {'-'*15} {'-'*30}")
        
        # Define display order
        category_order = ["AR", "RR", "Di", "Op", "Custom"]
        category_full_names = {"AR": "P-series (A)", "RR": "P-series (R)", "Di": "Device-direct", "Op": "Option codes", "Custom": "Custom/Special"}
        
        total_duration_seconds = 0
        for cat in category_order:
            if cat in category_times:
                info = category_times[cat]
                full_name = category_full_names.get(cat, cat)
                time_range = f"{info['start_time']} - {info['end_time']}"
                
                if info['is_overnight']:
                    print(f"{full_name:<20} {info['test_count']:<8} {info['duration']:<15} (Overnight)")
                    
                    daily_breakdown_str = ""
                    for day in info['daily_breakdown']:
                        print(f"  └─ {day['date']}: {day['test_count']} tests, {day['duration']} ({day['start_time']} - {day['end_time']})")
                        if daily_breakdown_str:
                            daily_breakdown_str += " | "
                        daily_breakdown_str += f"{day['date']}: {day['duration']} ({day['test_count']} tests)"
                    
                    test_time_output.append([full_name, info['test_count'], info['start_time'], info['end_time'], info['duration'], daily_breakdown_str])
                    
                    for day in info['daily_breakdown']:
                        test_time_output.append([f"  {full_name} ({day['date']})", day['test_count'], day['start_time'], day['end_time'], day['duration'], ""])
                else:
                    print(f"{full_name:<20} {info['test_count']:<8} {info['duration']:<15} {time_range}")
                    test_time_output.append([full_name, info['test_count'], info['start_time'], info['end_time'], info['duration'], ""])
                
                total_duration_seconds += info['duration_seconds']
        
        # Calculate and display total duration
        total_hours, total_remainder = divmod(total_duration_seconds, 3600)
        total_minutes, total_seconds = divmod(total_remainder, 60)
        total_duration_str = f"{total_hours}h {total_minutes}m {total_seconds}s"
        print(f"{'-'*20} {'-'*8} {'-'*15} {'-'*30}")
        print(f"{'TOTAL':<20} {sum(info['test_count'] for info in category_times.values()):<8} {total_duration_str:<15}")
        test_time_output.append(["TOTAL", sum(info['test_count'] for info in category_times.values()), "", "", total_duration_str, ""])
        
        print(f"{'='*60}\n")
        test_time_output.append([])  # Spacer
    
    # Calculate and display overall test time (across all categories)
    overall_time = calculate_overall_test_time(data, parse_date)
    overall_time_output = []
    
    if overall_time:
        print(f"\n{'='*60}")
        header_msg = "--- Overall Test Duration (All Tests Combined) ---"
        print(f"{header_msg}")
        print(f"{'='*60}")
        overall_time_output.append([header_msg])
        overall_time_output.append(["Date", "Test Count", "Start Time", "End Time", "Duration"])
        
        # Print formatted table header
        print(f"\n{'Date':<15} {'Tests':<8} {'Duration':<15} {'Time Range':<30}")
        print(f"{'-'*15} {'-'*8} {'-'*15} {'-'*30}")
        
        # Display daily breakdown
        for day in overall_time['daily_breakdown']:
            time_range = f"{day['start_time']} - {day['end_time']}"
            print(f"{day['date']:<15} {day['test_count']:<8} {day['duration']:<15} {time_range}")
            overall_time_output.append([day['date'], day['test_count'], day['start_time'], day['end_time'], day['duration']])
        
        # Display total
        print(f"{'-'*15} {'-'*8} {'-'*15} {'-'*30}")
        if overall_time['num_days'] > 1:
            print(f"{'TOTAL':<15} {overall_time['test_count']:<8} {overall_time['duration']:<15} ({overall_time['num_days']} days)")
        else:
            print(f"{'TOTAL':<15} {overall_time['test_count']:<8} {overall_time['duration']:<15}")
        overall_time_output.append(["TOTAL", overall_time['test_count'], overall_time['start_time'], overall_time['end_time'], overall_time['duration']])
        
        print(f"{'='*60}\n")
        overall_time_output.append([])  # Spacer
    
    # Check for Shared Speakerphone requirements if applicable
    validation_output = []
    if is_shared_speakerphone:
        print(f"\n{'='*80}")
        header_msg = "--- Shared Space Speakerphone Validation ---"
        print(f"{header_msg}")
        print(f"{'='*80}")
        validation_output.append([header_msg])
        
        # Display Equipment Settings from SmdSettings
        if smd_settings:
            print(f"\n--- Equipment Settings (from SmdSettings) ---")
            validation_output.append([""])
            validation_output.append(["--- Equipment Settings ---"])
            
            # Display labCORE information
            if smd_settings['labCORE']:
                print(f"\n{'labCORE Information:'}")
                validation_output.append(["labCORE Information"])
                validation_output.append(["File", "Serial", "Firmware", "Nickname"])
                
                print(f"  {'File':<40} {'Serial':<20} {'Firmware':<15} {'Nickname':<30}")
                print(f"  {'-'*40} {'-'*20} {'-'*15} {'-'*30}")
                
                for lc in smd_settings['labCORE']:
                    file_name = lc.get('file', '')[:38]
                    serial = lc.get('serial', 'N/A')
                    firmware = lc.get('firmware', 'N/A')
                    nickname = lc.get('nickname', 'N/A')
                    print(f"  {file_name:<40} {serial:<20} {firmware:<15} {nickname:<30}")
                    validation_output.append([lc.get('file', ''), serial, firmware, nickname])
            
            # Display HATS information
            if smd_settings['HATS']:
                print(f"\n{'HATS/HMS Information:'}")
                validation_output.append([""])
                validation_output.append(["HATS/HMS Information"])
                validation_output.append(["File", "Serial", "Equalization", "Pinna Type"])
                
                print(f"  {'File':<40} {'Serial':<15} {'Equalization':<15} {'Pinna Type':<25}")
                print(f"  {'-'*40} {'-'*15} {'-'*15} {'-'*25}")
                
                for hats in smd_settings['HATS']:
                    file_name = hats.get('file', '')[:38]
                    serial = hats.get('serial', 'N/A')
                    equalization = hats.get('equalization', 'N/A')
                    pinna = hats.get('pinna', 'N/A')
                    print(f"  {file_name:<40} {serial:<15} {equalization:<15} {pinna:<25}")
                    validation_output.append([hats.get('file', ''), serial, equalization, pinna])
            
            # Display BEQ Settings (especially for P05R/P10R with DF)
            if smd_settings['BEQ']:
                print(f"\n{'BEQ Equalization Settings:'}")
                validation_output.append([""])
                validation_output.append(["BEQ Equalization Settings"])
                validation_output.append(["File", "Test Code", "Equalization", "BEQ Setting", "Has DF"])
                
                print(f"  {'File':<40} {'Test Code':<10} {'Equalization':<25} {'Has DF':<10}")
                print(f"  {'-'*40} {'-'*10} {'-'*25} {'-'*10}")
                
                for beq in smd_settings['BEQ']:
                    file_name = beq.get('file', '')[:38]
                    test_code = beq.get('test_code', '')
                    equalization = beq.get('equalization', beq.get('beq_setting', 'N/A'))[:23]
                    has_df = "Yes" if beq.get('has_df', False) else "No"
                    
                    # Highlight P05R/P10R with DF
                    if test_code in ['P05R', 'P10R'] and beq.get('has_df', False):
                        print(f"  {file_name:<40} {test_code:<10} {equalization:<25} {has_df:<10} *** DF for {test_code}")
                    else:
                        print(f"  {file_name:<40} {test_code:<10} {equalization:<25} {has_df:<10}")
                    
                    validation_output.append([beq.get('file', ''), test_code, beq.get('equalization', beq.get('beq_setting', '')), beq.get('beq_setting', ''), has_df])
            
            print("")
        
        # ACQUA & Teams Database Information
        print(f"\n--- ACQUA & Teams Database Information ---")
        validation_output.append([""])
        validation_output.append(["--- ACQUA & Teams Database Information ---"])
        validation_output.append(["File", "ACQUA Version", "Database Version"])
        
        print(f"  {'File':<60} {'ACQUA Version':<20} {'Database Version':<40}")
        print(f"  {'-'*60} {'-'*20} {'-'*40}")
        
        for info in acqua_db_info:
            file_name = info.get('file', '')[:58]
            acqua_ver = info.get('acqua_version', 'Not Found')
            db_ver = info.get('database_version', 'Not Found')
            print(f"  {file_name:<60} {acqua_ver:<20} {db_ver:<40}")
            validation_output.append([info.get('file', ''), acqua_ver, db_ver])
        
        print("")
        
        # Test case validation
        print(f"\n--- Test Case Validation ---")
        validation_output.append([""])
        validation_output.append(["--- Test Case Validation ---"])

        # Add minimal subset test cases output
        minimal_subset = [
            "P01A", "P02A", "P03A", "P04A", "P09A", "P12A", "P13A", "P14A", "P21A", "P25A", "P26A", "P27A", "P01D", "P01R", "P02R", "P10R", "P11R", "P12R"
        ]
        minimal_subset_str = "Minimal subset test cases are: " + ", ".join(minimal_subset)
        print(f"\n{minimal_subset_str}\n")
        validation_output.append([minimal_subset_str])
        
        found_code_ids = set()
        for row in data:
            # row format: [CodeID, Category, SmdTitle_Times, SmdTitle, SmdDate]
            found_code_ids.add(row[0]) 

        print(f"\n{'Category':<15} {'Status':<50}")
        print(f"{'-'*15} {'-'*50}")
        
        for cat, required_codes in REQUIRED_SHARED_SPEAKERPHONE.items():
            missing = required_codes - found_code_ids
            if not missing:
                status = "✓ All required test cases included"
                print(f"{cat:<15} {status}")
                validation_output.append([f"Category {cat}: All required test cases included"])
            else:
                status = f"✗ Missing: {', '.join(sorted(missing))}"
                print(f"{cat:<15} {status}")
                validation_output.append([f"Category {cat}: Missing - {', '.join(sorted(missing))}"])
        
        print(f"{'='*80}\n")
        validation_output.append([])  # Spacer
    
    # Display and prepare 54dB noise scenario results
    noise_output = []
    if noise_54db_results:
        print(f"\n{'='*100}")
        header_msg = "--- 54dB Noise Scenario Results (NS ON vs NS OFF) ---"
        print(f"{header_msg}")
        print(f"{'='*100}")
        noise_output.append([header_msg])
        
        # Header row
        noise_header = ["Device", "Lab", "Report Time", "NS_Setting", "SMOS_2ndTalker", "NMOS_2ndTalker", "GMOS_2ndTalker", "SMOS_BGN", "NMOS_BGN", "GMOS_BGN"]
        noise_output.append(noise_header)
        
        # Print formatted table header
        print(f"\n{'Device':<100} {'Lab':<6} {'Date':<12} {'NS':<8} {'SMOS(2nd)':<10} {'NMOS(2nd)':<10} {'GMOS(2nd)':<10} {'SMOS(BGN)':<10} {'NMOS(BGN)':<10} {'GMOS(BGN)':<10}")
        print(f"{'-'*100} {'-'*6} {'-'*12} {'-'*8} {'-'*10} {'-'*10} {'-'*10} {'-'*10} {'-'*10} {'-'*10}")
        
        for result in noise_54db_results:
            row = [
                result['Device'],
                result['Lab'],
                result['Report_Time'],
                result['NS_Setting'],
                result['SMOS_2ndTalker'],
                result['NMOS_2ndTalker'],
                result['GMOS_2ndTalker'],
                result['SMOS_BGN'],
                result['NMOS_BGN'],
                result['GMOS_BGN']
            ]
            # Console output with longer device name (up to 100 chars)
            device = str(result['Device'])[:98] if len(str(result['Device'])) > 98 else result['Device']
            report_date = str(result['Report_Time'])[:12] if result['Report_Time'] else ''
            print(f"{device:<100} {result['Lab']:<6} {report_date:<12} {result['NS_Setting']:<8} {result['SMOS_2ndTalker']:<10} {result['NMOS_2ndTalker']:<10} {result['GMOS_2ndTalker']:<10} {result['SMOS_BGN']:<10} {result['NMOS_BGN']:<10} {result['GMOS_BGN']:<10}")
            noise_output.append(row)
        
        print(f"{'='*100}\n")
        noise_output.append([])  # Spacer
    
    # Display "Not OK" status rows from Status Overview table
    status_output = []
    if not_ok_rows:
        header_msg = "--- Status Overview: 'Not OK' Entries ---"
        print(f"\n{'='*80}")
        print(f"{header_msg}")
        print(f"{'='*80}")
        status_output.append([header_msg])
        
        # Header row (now includes Limit column)
        status_header = ["FileName", "SMD", "Status", "Single Value Description", "Single Value", "Limit"]
        status_output.append(status_header)
        
        # Print formatted table header
        print(f"\n{'No.':<5} {'SMD':<55} {'Status':<10} {'Single Value':<15} {'Limit':<20}")
        print(f"{'-'*5} {'-'*55} {'-'*10} {'-'*15} {'-'*20}")
        
        # Group by filename for better readability
        current_file = ""
        entry_num = 0
        for entry in not_ok_rows:
            # Print filename header when it changes
            if entry['FileName'] != current_file:
                current_file = entry['FileName']
                print(f"\n[{current_file}]")
            
            entry_num += 1
            smd_short = entry['SMD'][:53] + '..' if len(entry['SMD']) > 55 else entry['SMD']
            single_val = str(entry['SingleValue'])[:13] if entry['SingleValue'] else ''
            limit_val = str(entry.get('Limit', ''))[:18] if entry.get('Limit') else ''
            
            print(f"{entry_num:<5} {smd_short:<55} {entry['Status']:<10} {single_val:<15} {limit_val:<20}")
            
            # Full row for CSV
            row = [
                entry['FileName'],
                entry['SMD'],
                entry['Status'],
                entry['SingleValueDescription'],
                entry['SingleValue'],
                entry.get('Limit', '')
            ]
            status_output.append(row)
        
        print(f"\n{'='*80}")
        print(f"Total 'Not OK' entries: {len(not_ok_rows)}")
        print(f"{'='*80}\n")
        status_output.append([f"Total 'Not OK' entries: {len(not_ok_rows)}"])
        status_output.append([])  # Spacer
    elif all_status_rows:
        ok_msg = f"Status Overview: All {len(all_status_rows)} entries are OK!"
        print(f"\n{'='*80}")
        print(f"{ok_msg}")
        print(f"{'='*80}\n")
        status_output.append(["--- Status Overview ---"])
        status_output.append([ok_msg])
        status_output.append([])  # Spacer

    # Display and prepare double talk performance results
    double_talk_output = []
    if double_talk_results:
        print(f"\n{'='*100}")
        header_msg = "--- Double Talk Performance (Attenuation during double talk [dB]) ---"
        print(f"{header_msg}")
        print(f"{'='*100}")
        double_talk_output.append([header_msg])
        
        # Header row
        dt_header = ["FileName", "SMD", "Status", "Description", "Single Value (dB)"]
        double_talk_output.append(dt_header)
        
        # Print formatted table header
        print(f"\n{'No.':<5} {'SMD':<60} {'Status':<10} {'Value (dB)':<15}")
        print(f"{'-'*5} {'-'*60} {'-'*10} {'-'*15}")
        
        # Group by filename for better readability
        current_file = ""
        entry_num = 0
        for entry in double_talk_results:
            # Print filename header when it changes
            if entry['FileName'] != current_file:
                current_file = entry['FileName']
                print(f"\n[{current_file}]")
            
            entry_num += 1
            smd_short = entry['SMD'][:58] + '..' if len(entry['SMD']) > 60 else entry['SMD']
            single_val = str(entry['SingleValue'])[:13] if entry['SingleValue'] else ''
            
            print(f"{entry_num:<5} {smd_short:<60} {entry['Status']:<10} {single_val:<15}")
            
            # Full row for CSV
            row = [
                entry['FileName'],
                entry['SMD'],
                entry['Status'],
                entry['Description'],
                entry['SingleValue']
            ]
            double_talk_output.append(row)
        
        print(f"\n{'='*100}")
        print(f"Total double talk entries: {len(double_talk_results)}")
        print(f"{'='*100}\n")
        double_talk_output.append([f"Total double talk entries: {len(double_talk_results)}"])
        double_talk_output.append([])  # Spacer

    # Generate Output File
    # Use the directory of the first selected file
    output_dir = os.path.dirname(file_paths[0])
    output_filename = os.path.join(output_dir, "Smd_Report_Output.csv")
    
    try:
        with open(output_filename, mode='w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file)
            
            # Write test time results first
            if test_time_output:
                writer.writerows(test_time_output)
            
            # Write overall test time results
            if overall_time_output:
                writer.writerows(overall_time_output)
            
            # Write validation results
            if validation_output:
                writer.writerows(validation_output)
            
            # Write 54dB noise scenario results
            if noise_output:
                writer.writerows(noise_output)
            
            # Write Status Overview "Not OK" entries
            if status_output:
                writer.writerows(status_output)
            
            # Write double talk performance results
            if double_talk_output:
                writer.writerows(double_talk_output)

            # Writing headers as requested
            writer.writerow(["CodeID", "Category", "SmdTitle_Times", "SmdTitle", "SmdDate", "FilePath"])
            writer.writerows(data)
            
        print(f"\n{'='*60}")
        print(f"✓ SUCCESS - Extracted {len(data)} test entries")
        print(f"  Output saved to: {os.path.abspath(output_filename)}")
        print(f"{'='*60}")
        
    except PermissionError:
        print(f"Error: Could not write to {output_filename}. Is the file open in Excel?")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description=__description__,
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument(
        '--version', '-v',
        action='version',
        version=f'%(prog)s {__version__}'
    )
    
    args = parser.parse_args()
    main()
